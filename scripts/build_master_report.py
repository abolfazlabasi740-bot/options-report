from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS = Path(
    r"C:\Users\ab.abasi\.gapcode\plugins\cache\gapgpt\documents"
    r"\26.722.10000\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))

from fa_docx import persian_digits, set_font, set_rtl_paragraph, set_rtl_styles, set_rtl_table
from table_geometry import apply_table_geometry, column_widths_from_weights

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


FONT = "Vazirmatn"
NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
WHITE = "FFFFFF"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WARNING_FILL = "FFF4CE"
GOOD_FILL = "EAF4EA"
RISK_FILL = "FCE8E6"
BORDER = "AAB7C4"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def latest_file(directory: Path, pattern: str) -> Path:
    files = list(directory.glob(pattern))
    if not files:
        raise FileNotFoundError(f"No files matched {pattern} in {directory}")
    return max(files, key=lambda p: p.stat().st_mtime)


def fmt_int(value) -> str:
    return persian_digits(f"{int(round(float(value))):,}")


def fmt_num(value, digits=2) -> str:
    return persian_digits(f"{float(value):.{digits}f}")


def fmt_pct(value, digits=3) -> str:
    return persian_digits(f"{float(value) * 100:.{digits}f}") + "٪"


def localized_bool(value: bool) -> str:
    return "بله" if value else "خیر"


def grouped_hash(value: str) -> str:
    return "\n".join(value[i : i + 16] for i in range(0, len(value), 16))


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def configure_styles(doc: Document):
    set_rtl_styles(doc, FONT)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_specs = {
        "Title": (24, NAVY, 0, 8),
        "Subtitle": (12, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, FONT, 9)


def set_section_furniture(section, run_id: str):
    header = section.header.paragraphs[0]
    header.text = "گزارش اجرایی تحلیل اختیار معامله | Baseline رسمی V3"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(header, FONT, 8.5)
    for run in header.runs:
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Run ID: {run_id} | صفحه ")
    set_font(run, FONT, 9)
    add_page_field(footer)
    set_rtl_paragraph(footer, FONT, 9)
    for item in footer.runs:
        item.font.color.rgb = MUTED


def add_rtl(doc, text="", style=None, bold=False, color=None, size=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_font(run, FONT, size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
    if align is not None:
        paragraph.alignment = align
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(paragraph, FONT, size)
    return paragraph


def add_heading(doc, text, level=1):
    return add_rtl(doc, text, style=f"Heading {level}")


def write_cell(cell, text, *, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(str(text))
    set_font(run, FONT, size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    set_rtl_paragraph(paragraph, FONT, size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(
    doc,
    headers,
    rows,
    weights,
    *,
    total_width=9360,
    font_size=9.5,
    header_fill=HEADER_FILL,
    alignments=None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_rtl_table(table, FONT)
    for idx, header in enumerate(headers):
        write_cell(table.rows[0].cells[idx], header, bold=True, size=font_size)
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_borders(table.rows[0].cells[idx])
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            alignment = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(row.cells[idx], value, size=font_size, align=alignment)
            set_cell_borders(row.cells[idx])

    widths = column_widths_from_weights(weights, total_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=total_width,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_rtl_table(table, FONT)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_status_callout(doc, title, text, fill=GOOD_FILL):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), fill)
        borders.append(border)
    p_pr.append(borders)
    title_run = paragraph.add_run(title + ": ")
    set_font(title_run, FONT, 10.5)
    title_run.bold = True
    text_run = paragraph.add_run(text)
    set_font(text_run, FONT, 10.5)
    set_rtl_paragraph(paragraph, FONT, 10.5)
    return paragraph


def build_report(project_root: Path, output_path: Path | None = None) -> Path:
    run_path = latest_file(project_root / "data" / "snapshots" / "runs", "*.json")
    run = load_json(run_path)
    run_id = run["RunId"]
    audit_path = project_root / "logs" / "audit" / f"{run_id}.json"
    learning_path = project_root / "library" / "reports" / f"{run_id}.json"
    raw_experience_path = project_root / "library" / "raw_experience" / f"EXP_{run_id}.json"
    markdown_path = project_root / "reports" / f"{run_id}_options_report.md"
    config_path = project_root / "configs" / "project.json"
    instruction_path = project_root / "reference" / "Master Project Book.docx"
    if not instruction_path.exists():
        # Backward-compatible path for the original local layout.
        instruction_path = project_root.parent / "master project book" / "Master Project Book.docx"

    required = [
        audit_path,
        learning_path,
        raw_experience_path,
        markdown_path,
        config_path,
        instruction_path,
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required Master Project Book evidence:\n" + "\n".join(missing))

    audit = load_json(audit_path)
    learning = load_json(learning_path)
    config = load_json(config_path)
    top = run["Top"]
    scores = {item["Symbol"]: item for item in audit["Scores"]}

    if audit["RunId"] != run_id or learning["RunId"] != run_id:
        raise ValueError("Run, audit, and learning artifacts do not share one Run ID.")
    if Path(audit["Input"]["Workbook"]).name != run["FileName"]:
        raise ValueError("Audit input workbook does not match the run snapshot.")
    if audit["Input"]["Sha256"] != run["FileSha256"]:
        raise ValueError("Audit and run workbook hashes do not match.")

    if output_path is None:
        output_path = project_root / "reports" / f"گزارش_اجرایی_Master_V3_{run_id}.docx"

    missing_iv = sum(1 for item in top if item.get("IV") is None)
    wide_spread = sum(1 for item in top if "WIDE_SPREAD" in item.get("Flags", []))
    avg_leverage = sum(item["Leverage"] for item in top) / len(top)
    avg_days = sum(item["RemainingDays"] for item in top) / len(top)
    avg_penalty = sum(item["RiskPenalty"] for item in top) / len(top)
    total_volume = sum(item["Volume"] for item in top)
    total_value = sum(item["TradeValue"] for item in top)
    rank_changes = learning["Comparison"].get("RankChanges", [])
    changed_ranks = sum(1 for item in rank_changes if item.get("Change") != 0)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    set_section_furniture(doc.sections[0], run_id)

    title = add_rtl(
        doc,
        "گزارش اجرایی تحلیل اختیار معامله",
        style="Title",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    title.paragraph_format.space_before = Pt(22)
    add_rtl(
        doc,
        "مطابق الزامات Evidence، Lineage و Reporting در Master Project Book",
        style="Subtitle",
    )
    add_rtl(
        doc,
        f"اجرای تازه Pipeline رسمی | Run ID: {run_id}",
        bold=True,
        color=DARK_BLUE,
        size=11,
    )

    add_status_callout(
        doc,
        "وضعیت اجرای تازه",
        (
            f"Pipeline با موفقیت کامل شد؛ {fmt_int(run['SourceRows'])} ردیف دریافت، "
            f"{fmt_int(run['ValidRows'])} قرارداد معتبر و {fmt_int(run['RemovedRows'])} "
            "ردیف حذف‌شده ثبت شد."
        ),
        GOOD_FILL,
    )
    add_status_callout(
        doc,
        "مرز نسخه",
        (
            f"پروتکل اجرایی این Run {run['Protocol']} و Baseline رسمی V3 است؛ "
            "V4 فقط Candidate است و Specification اجرایی تأییدشده ندارد."
        ),
        WARNING_FILL,
    )
    add_status_callout(
        doc,
        "ماهیت تصمیم",
        (
            "خروجی فقط رتبه‌بندی کیفیت قرارداد است؛ سیگنال خرید یا فروش تولید نشده، "
            "جهت بازار ارائه نشده و آستانه‌های طبقه‌بندی تأیید نشده‌اند."
        ),
        RISK_FILL,
    )

    add_heading(doc, "۱. شناسنامه اجرا و Evidence", 1)
    metadata_rows = [
        ("Run ID", run_id),
        ("زمان پردازش", run["Timestamp"]),
        ("فایل داده جاری", run["FileName"]),
        ("پروتکل امتیازدهی", run["Protocol"]),
        ("نسخه پروژه", run["Version"]),
        ("وضعیت Pipeline", "COMPLETED"),
        ("وضعیت Audit", audit["Outcome"]),
        ("وضعیت Learning", learning["ValidationStatus"]),
    ]
    add_table(
        doc,
        ["مشخصه", "مقدار"],
        metadata_rows,
        [2, 5],
        font_size=9.5,
        alignments=[WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
    )

    add_heading(doc, "۲. کیفیت داده و کنترل ورودی", 1)
    quality_rows = [
        ("ردیف‌های منبع", fmt_int(run["SourceRows"]), "ثبت‌شده در Run Snapshot"),
        ("قراردادهای معتبر", fmt_int(run["ValidRows"]), "عبور از فیلترهای سخت"),
        ("ردیف‌های حذف‌شده", fmt_int(run["RemovedRows"]), "هیچ حذف ثبت نشده است"),
        ("قرارداد Call", fmt_int(run["Calls"]), "نوع تشخیص‌داده‌شده"),
        ("قرارداد Put", fmt_int(run["Puts"]), "در داده جاری موجود نیست"),
        ("نوع نامشخص", fmt_int(run["UnknownType"]), "موردی ثبت نشده است"),
        ("IV ناموجود در ۱۵ رتبه برتر", fmt_int(missing_iv), "وزن در بلوک مربوط بازتوزیع شده است"),
        ("پرچم شکاف قیمتی باز", fmt_int(wide_spread), "WIDE_SPREAD در Audit"),
    ]
    add_table(
        doc,
        ["شاخص", "مقدار", "وضعیت / Evidence"],
        quality_rows,
        [2.5, 1.2, 3.3],
        font_size=9.3,
        alignments=[
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )

    removed_rows = [
        (reason, fmt_int(count)) for reason, count in run["RemovedByReason"].items()
    ]
    add_heading(doc, "دلایل حذف رکورد", 2)
    add_table(
        doc,
        ["کد دلیل", "تعداد"],
        removed_rows,
        [4.7, 1.3],
        total_width=8640,
        font_size=9.2,
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    set_section_furniture(landscape, run_id)

    add_heading(doc, "۳. رتبه‌بندی قراردادهای برتر", 1)
    add_rtl(
        doc,
        "این جدول مستقیماً از Top موجود در Run Snapshot همین اجرا ساخته شده است.",
        size=9.5,
        color=MUTED,
    )
    ranking_rows = []
    for rank, item in enumerate(top, 1):
        ranking_rows.append(
            (
                fmt_int(rank),
                item["Symbol"],
                fmt_int(item["Strike"]),
                fmt_int(item["Last"]),
                fmt_int(item["Breakeven"]),
                fmt_int(item["Underlying"]),
                fmt_num(item["Leverage"]),
                fmt_pct(item["BEDistance"]),
                item["Expiration"],
                fmt_int(item["RemainingDays"]),
                fmt_num(item["FinalScore"]),
            )
        )
    add_table(
        doc,
        [
            "رتبه",
            "نماد",
            "اعمال",
            "آخرین",
            "سربه‌سر",
            "پایه",
            "اهرم",
            "فاصله سربه‌سر",
            "سررسید",
            "روز",
            "امتیاز",
        ],
        ranking_rows,
        [0.55, 1.15, 0.8, 0.8, 0.85, 0.85, 0.6, 1.15, 1.0, 0.55, 0.7],
        total_width=12960,
        font_size=8.2,
    )

    add_heading(doc, "۴. اجزای امتیاز و جریمه ریسک", 1)
    score_rows = []
    for rank, item in enumerate(top, 1):
        score = scores[item["Symbol"]]
        blocks = score["Blocks"]
        score_rows.append(
            (
                fmt_int(rank),
                item["Symbol"],
                fmt_num(blocks["Liquidity"]),
                fmt_num(blocks["Valuation"]),
                fmt_num(blocks["Payoff"]),
                fmt_num(blocks["Time"]),
                fmt_num(blocks["Greeks"]),
                fmt_num(blocks["Market"]),
                fmt_num(score["Base"]),
                fmt_num(score["Penalty"]),
                fmt_num(score["Final"]),
            )
        )
    add_table(
        doc,
        [
            "رتبه",
            "نماد",
            "نقدشوندگی",
            "ارزش‌گذاری",
            "بازده",
            "زمان",
            "یونانی‌ها",
            "بازار",
            "پایه",
            "جریمه",
            "نهایی",
        ],
        score_rows,
        [0.5, 1.1, 0.9, 0.95, 0.8, 0.7, 0.9, 0.75, 0.75, 0.75, 0.75],
        total_width=12960,
        font_size=8.1,
    )

    add_heading(doc, "۵. کیفیت داده رتبه‌های برتر", 1)
    missing_rows = []
    for rank, item in enumerate(top, 1):
        missing = "، ".join(item.get("MissingData", [])) or "ندارد"
        flags = "، ".join(item.get("Flags", [])) or "ندارد"
        missing_rows.append(
            (
                fmt_int(rank),
                item["Symbol"],
                fmt_num(item["DataConfidence"] * 100, 0) + "٪",
                missing,
                flags,
                item["DecisionType"],
                item["Classification"],
            )
        )
    add_table(
        doc,
        ["رتبه", "نماد", "اطمینان داده", "داده مفقود", "پرچم‌ها", "نوع تصمیم", "طبقه‌بندی"],
        missing_rows,
        [0.5, 1.0, 0.85, 1.1, 1.6, 1.7, 1.7],
        total_width=12960,
        font_size=8.0,
    )

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(portrait, landscape=False)
    set_section_furniture(portrait, run_id)

    add_heading(doc, "۶. خلاصه مدیریتی", 1)
    summary_rows = [
        ("میانگین اهرم ۱۵ قرارداد برتر", fmt_num(avg_leverage), "برابر"),
        ("میانگین روز باقی‌مانده", fmt_num(avg_days, 1), "روز"),
        ("حجم تجمیعی معاملات", fmt_int(total_volume), "قرارداد"),
        ("ارزش تجمیعی معاملات", fmt_int(total_value), "واحد پولی فایل"),
        ("میانگین جریمه ریسک", fmt_num(avg_penalty), "امتیاز"),
        ("قراردادهای مشترک با Run قبلی", fmt_int(learning["Comparison"].get("CommonContracts", 0)), "قرارداد"),
        ("رتبه‌های تغییرکرده", fmt_int(changed_ranks), "از میان قراردادهای مشترک"),
    ]
    add_table(
        doc,
        ["شاخص", "مقدار", "واحد / توضیح"],
        summary_rows,
        [3.1, 1.3, 2.6],
        font_size=9.5,
        alignments=[
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )

    add_heading(doc, "۷. وضعیت Decision، Learning و Knowledge", 1)
    governance_rows = [
        ("خروجی تصمیم", audit["Decision"]["Scenario"], "رتبه‌بندی کیفیت بدون جهت بازار"),
        (
            "سیگنال خرید",
            localized_bool(audit["Decision"]["GeneratesBuySignal"]),
            "تولید نشده است",
        ),
        (
            "سیگنال فروش",
            localized_bool(audit["Decision"]["GeneratesSellSignal"]),
            "تولید نشده است",
        ),
        (
            "طبقه‌بندی عددی",
            audit["Decision"]["ClassificationStatus"],
            "آستانه‌ها تأیید نشده‌اند",
        ),
        (
            "مقایسه با Run قبلی",
            learning["Comparison"]["Status"],
            f"Run قبلی: {learning['Comparison'].get('PreviousRunId', 'داده موجود نیست')}",
        ),
        (
            "اعتبارسنجی نتیجه",
            learning["Comparison"]["OutcomeValidation"],
            "پیش‌بینی جهت‌دار برای ارزیابی وجود ندارد",
        ),
        (
            "تغییر خودکار Rule",
            localized_bool(learning["AutomaticRuleChangeAllowed"]),
            "طبق Master Project Book مجاز نیست",
        ),
        ("وضعیت Learning", learning["ValidationStatus"], "در حال گردآوری Evidence"),
        ("Outcome در Audit", audit["Outcome"], "منتظر داده واقعی بعدی"),
    ]
    add_table(
        doc,
        ["لایه", "وضعیت ثبت‌شده", "تفسیر محدود و قابل ردیابی"],
        governance_rows,
        [2.1, 2.3, 2.6],
        font_size=9.0,
        alignments=[
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )

    add_heading(doc, "۸. کاتالوگ محاسبات مشتق‌شده", 1)
    feature_rows = [
        (
            item.get("Name", item.get("name", "")),
            item.get("Source", item.get("source", "")),
            item.get("Formula", item.get("formula", "")),
            item.get("Unit", item.get("unit", "")),
        )
        for item in run["FeatureCatalog"]
    ]
    add_table(
        doc,
        ["Feature", "منبع", "فرمول ثبت‌شده", "واحد"],
        feature_rows,
        [1.4, 1.1, 3.7, 0.8],
        font_size=8.8,
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
    )

    add_heading(doc, "۹. Lineage و Artifactهای مبنا", 1)
    lineage_rows = [
        ("Master Project Book", instruction_path.name, "مرجع رسمی حاکمیت، Baseline و گزارش‌گری"),
        ("Raw Evidence", f"data/raw/{run['FileName']}", "فایل داده جاری این Run"),
        ("Run Snapshot", f"data/snapshots/runs/{run_id}.json", "ورودی، Featureها و Top"),
        ("Audit Record", f"logs/audit/{run_id}.json", "Lineage، Score و Decision"),
        ("Learning Record", f"library/reports/{run_id}.json", "مقایسه Run و وضعیت یادگیری"),
        (
            "Raw Experience",
            f"library/raw_experience/EXP_{run_id}.json",
            "Evidence خام برای Knowledge",
        ),
        ("گزارش موتور", f"reports/{run_id}_options_report.md", "خروجی deterministic موتور Reporting"),
        ("Config", "configs/project.json", "نسخه و وزن‌های اجرایی"),
    ]
    add_table(
        doc,
        ["نوع Artifact", "مسیر نسبی", "نقش در Lineage"],
        lineage_rows,
        [1.6, 3.3, 2.1],
        font_size=8.8,
        alignments=[
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )

    add_heading(doc, "۱۰. اطلاعات یکپارچگی", 1)
    integrity_rows = [
        ("SHA-256 فایل داده", grouped_hash(run["FileSha256"])),
        ("SHA-256 تنظیمات", grouped_hash(run["ConfigSha256"])),
        ("نسخه ثبت‌شده در Config", config["version"]),
        ("پروتکل ثبت‌شده در Config", config["protocol"]),
    ]
    add_table(
        doc,
        ["مشخصه", "مقدار"],
        integrity_rows,
        [2.0, 5.0],
        font_size=8.5,
        alignments=[WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(doc, "۱۱. محدودیت‌ها و وضعیت تأیید", 1)
    limitations = [
        "این گزارش هیچ داده، وزن، آستانه، فرمول یا نتیجه‌ای خارج از Artifactهای همین Run ایجاد نکرده است.",
        "فایل داده جاری مبنای تحلیل است؛ فایل قدیمی جایگزین آن نشده است.",
        "Scoring اجرایی Run با پروتکل V3 انجام شده و V4 فقط Candidate است.",
        f"نوسان ضمنی در {missing_iv} قرارداد از {len(top)} قرارداد برتر موجود نیست و این وضعیت در Audit ثبت شده است.",
        f"جامعه معتبر شامل {run['Calls']} قرارداد Call و {run['Puts']} قرارداد Put است.",
        f"Risk با وضعیت {run.get('RiskAlignmentStatus', 'NOT_RECORDED')} ثبت شده و آستانه‌های پله‌ای تأییدنشده حدس زده نشده‌اند.",
        "خروجی صرفاً تصمیم‌یار و رتبه‌بندی کیفیت است و توصیه خرید یا فروش محسوب نمی‌شود.",
        "Outcome فعلی PENDING_NEXT_REAL_DATA است و Production Readiness کل سامانه از این گزارش نتیجه‌گیری نمی‌شود.",
    ]
    for text in limitations:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run_text = paragraph.add_run(text)
        set_font(run_text, FONT, 10.5)
        set_rtl_paragraph(paragraph, FONT, 10.5)

    add_status_callout(
        doc,
        "نتیجه نهایی",
        (
            "اجرای Pipeline موفق، Lineage داخلی کامل و گزارش قابل بازتولید است؛ "
            "اما سیگنال معاملاتی، جهت بازار و تأیید Production از این Run استخراج نشده است."
        ),
        LIGHT_FILL,
    )

    doc.core_properties.title = "گزارش اجرایی تحلیل اختیار معامله - Master Baseline V3"
    doc.core_properties.subject = f"Run ID {run_id}"
    doc.core_properties.author = "OptimusAI Reporting Pipeline"
    doc.core_properties.keywords = "Options, Master Project Book, Evidence, Lineage, Audit, Scoring V3"
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Build the Master Project Book governed V3 baseline report.")
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    result = build_report(args.project_root.resolve(), args.output)
    print(result)


if __name__ == "__main__":
    main()
