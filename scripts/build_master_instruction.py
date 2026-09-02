from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SKILL_SCRIPTS = Path(
    r"C:\Users\ab.abasi\.gapcode\plugins\cache\gapgpt\documents"
    r"\26.722.10000\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))

from fa_docx import set_font, set_rtl_paragraph, set_rtl_styles, set_rtl_table
from table_geometry import apply_table_geometry

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


FONT = "Vazirmatn"
INK = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WARNING_FILL = "FFF4CE"
BORDER = "AAB7C4"


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_rtl(doc: Document, text: str, style: str | None = None, bold=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_font(run, FONT, size or 11)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    set_rtl_paragraph(p, FONT, size or 11)
    return p


def add_heading(doc: Document, text: str, level=1):
    return add_rtl(doc, text, style=f"Heading {level}", bold=True)


def add_table(doc: Document, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_rtl_table(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        set_cell_fill(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_font(r, FONT, 9.5)
        r.bold = True
        set_rtl_paragraph(p, FONT, 9.5)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = ""
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(value))
            set_font(r, FONT, 9.2)
            set_rtl_paragraph(p, FONT, 9.2)
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, title: str, text: str, fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_rtl_table(table)
    apply_table_geometry(
        table,
        [9360],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 120, "bottom": 120, "start": 120, "end": 120},
    )
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=120, bottom=120)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(title + ": ")
    set_font(r1, FONT, 10)
    r1.bold = True
    r2 = p.add_run(text)
    set_font(r2, FONT, 10)
    set_rtl_paragraph(p, FONT, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_rtl_styles(doc, FONT)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Options Analytics System | Master Baseline V3")
    set_font(run, FONT, 8.5)
    run.font.color.rgb = MUTED
    set_rtl_paragraph(header, FONT, 8.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("مرجع اجرایی نسخه ۳٫۲٫۰ — ۱۴۰۵/۰۵/۲۶")
    set_font(run, FONT, 8.5)
    run.font.color.rgb = MUTED
    set_rtl_paragraph(footer, FONT, 8.5)


def build(project_root: Path) -> Path:
    output = project_root / "docs" / "دستورالعمل_اجرایی_Master_Baseline_V3.docx"
    doc = Document()
    configure_document(doc)

    title = add_rtl(doc, "دستورالعمل اجرایی سامانه تحلیل اختیار معامله", bold=True, size=21, color=DARK_BLUE)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    subtitle = add_rtl(doc, "مبتنی بر Master Project Book — Baseline رسمی V3", size=12, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(16)

    add_table(
        doc,
        ["مشخصه", "مقدار"],
        [
            ("نسخه پروژه", "۳٫۲٫۰"),
            ("پروتکل Production", "PROTOCOL_OPTIONS_RANKING_V3"),
            ("وضعیت V4", "Candidate؛ غیرمجاز برای معرفی به‌عنوان Production"),
            ("مرجع", "master project book/Master Project Book.docx"),
            ("SHA-256 مرجع", "3cba9235181063f040f6df83f18e2b739395528f417d1179ec02eb2ce10f9268"),
            ("تاریخ هم‌ترازی", "۱۴۰۵/۰۵/۲۶ — ۲۰۲۶/۰۸/۱۷"),
        ],
        [2600, 6760],
    )

    add_callout(
        doc,
        "قاعده حاکم",
        "هیچ عدد، وزن، آستانه، فرمول، API یا نتیجه بدون Evidence ساخته نمی‌شود. داده یا منطق ناموجود باید صریحاً Missing یا تأییدنشده ثبت شود.",
        WARNING_FILL,
    )

    add_heading(doc, "۱. زنجیره رسمی اجرا", 1)
    add_rtl(
        doc,
        "Source Adapter ← Raw Evidence ← Validation ← Parsing/Normalization ← Financial ← Analytics/Features ← Scoring ← Risk ← Decision ← Strategy ← Reporting ← Audit ← Learning ← Knowledge",
    )
    add_rtl(
        doc,
        "ترتیب اجرایی در Pipeline از Data آغاز می‌شود و تمام Artifactهای Run باید تا Hash Manifest قابل بازیابی باشند.",
    )

    add_heading(doc, "۲. مدل امتیازدهی V3", 1)
    add_table(
        doc,
        ["بلوک", "وزن کل", "عوامل داخلی"],
        [
            ("Liquidity", "۲۰", "Trade Value ۳۵٪؛ Volume ۲۵٪؛ OI ۱۵٪؛ Spread ۱۵٪؛ Depth ۱۰٪"),
            ("Valuation", "۲۵", "BS Edge ۳۲٪؛ IV ۲۸٪؛ IV/HV ۲۰٪؛ Time Value ۲۰٪"),
            ("Payoff", "۱۸", "Break-even Distance ۵۵٫۵۶٪؛ Leverage ۲۷٫۷۸٪؛ Moneyness ۱۶٫۶۷٪"),
            ("Time", "۱۵", "Trading Days ۴۰٪؛ Calendar Days ۱۳٫۳۳٪؛ Theta ۴۶٫۶۷٪"),
            ("Greeks", "۱۲", "Delta ۳۳٫۳۳٪؛ Gamma ۲۵٪؛ Vega ۲۵٪؛ Rho ۱۶٫۶۷٪"),
            ("Market Structure", "۱۰", "Last vs Close ۴۰٪؛ Intraday Range ۳۰٪؛ Status ۳۰٪"),
        ],
        [1800, 1200, 6360],
    )
    add_rtl(
        doc,
        "Normalization مصوب V3 از Robust Percentile استفاده می‌کند. عامل Missing امتیاز ثابت نمی‌گیرد و وزن عوامل موجود داخل همان بلوک بازتوزیع می‌شود.",
    )

    add_heading(doc, "۳. Market Structure", 1)
    add_table(
        doc,
        ["عامل", "فرمول/روش اجرایی", "وضعیت"],
        [
            ("Last vs Close", "Abs(LastPercent-ClosePercent) و Percentile معکوس", "IMPLEMENTED_V3"),
            ("Intraday Range", "Abs(High-Low)/Last و Percentile معکوس", "IMPLEMENTED_V3"),
            ("Status", "فقط نگاشت عددی مصوب Config", "UNKNOWN_NOT_APPROVED"),
        ],
        [1900, 4860, 2600],
    )
    add_callout(
        doc,
        "Status",
        "مقادیر «در سود»، «بی‌تفاوت» و «در ضرر» وجود دارند، اما نگاشت عددی و جهت Call/Put مصوب نیست؛ بنابراین امتیاز Status ساخته نمی‌شود.",
        WARNING_FILL,
    )

    add_heading(doc, "۴. Risk Engine", 1)
    add_rtl(
        doc,
        "Master مدل هدف را بر مبنای پنج وضعیت Leverage، Trading Days، Theta، IV/HV و Break-even Distance و جریمه پله‌ای تعریف می‌کند.",
    )
    add_table(
        doc,
        ["تعداد وضعیت نامطلوب همزمان", "جریمه"],
        [("کمتر از سه", "۰"), ("سه", "۵"), ("چهار", "۱۰"), ("پنج", "۲۰")],
        [6500, 2860],
    )
    add_callout(
        doc,
        "Known Gap",
        "آستانه Percentile پنج وضعیت در Master موجود نیست. حالت STEP_COUNT آماده است ولی بدون Threshold معتبر فعال نمی‌شود؛ اجرای فعلی Transitional و صریحاً در Audit علامت‌گذاری می‌شود.",
        WARNING_FILL,
    )

    add_heading(doc, "۵. Decision، Strategy و Learning", 1)
    add_table(
        doc,
        ["حوزه", "قاعده اجرایی"],
        [
            ("Decision", "Data Quality → Liquidity → Contract Quality → Underlying Quality → Risk → Valuation → Scenario"),
            ("Signal", "هیچ سیگنال قطعی خرید یا فروش تولید نمی‌شود."),
            ("Strategy", "Rule ورود/خروج یا ساختار استراتژی بدون Evidence ساخته نمی‌شود."),
            ("Learning", "تجربه واقعی ثبت می‌شود؛ تغییر خودکار Rule و Weight در Production ممنوع است."),
            ("Knowledge", "تجربه ابتدا RAW_UNVALIDATED است و پس از تکرار و Validation قابل ارتقا است."),
        ],
        [2100, 7260],
    )

    add_heading(doc, "۶. Artifactهای اجباری هر Run", 1)
    add_table(
        doc,
        ["Artifact", "حداقل محتوا"],
        [
            ("Input Manifest", "نام، Source، Timestamp، SHA-256، Size و Schema Hash"),
            ("Run Manifest", "Run ID، Start/End، Protocol، Code و Config Version"),
            ("Feature Snapshot", "تمام Featureهای مصرف‌شده، Missing و Confidence"),
            ("Ranking Snapshot", "رتبه کامل، Blockها، Base، Penalty و Final"),
            ("Decision/Audit", "Gateها، Scenario، Flags، Evidence و Lineage"),
            ("Learning/Knowledge", "Outcome، Comparison و Raw Experience"),
            ("Hash Manifest", "Hash ورودی، Config، Master، کد و خروجی‌های Run"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "۷. دستورهای اجرایی", 1)
    commands = [
        "اجرای رسمی زمان‌بندی‌شده:",
        "powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\smart_money_project\\scripts\\run_scheduled.ps1 -ProjectRoot .\\smart_money_project",
        "تست حاکمیت Master:",
        ".\\smart_money_project\\tests\\master_baseline_test.ps1 -ProjectRoot .\\smart_money_project",
        "Smoke Test با فایل واقعی:",
        ".\\smart_money_project\\tests\\smoke_test.ps1 -WorkbookPath <xlsx> -ProjectRoot .\\smart_money_project",
    ]
    for index, text in enumerate(commands):
        p = add_rtl(doc, text, bold=(index % 2 == 0), size=9.5 if index % 2 else 10.5)
        if index % 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_heading(doc, "۸. Definition of Done", 1)
    add_rtl(
        doc,
        "Capability فقط وقتی Done است که Design، Implementation، Test، Evidence، Auditability و Versioning همزمان وجود داشته باشند. وجود کد به‌تنهایی Production Readiness را ثابت نمی‌کند.",
    )
    add_callout(
        doc,
        "مرز گزارش تازه",
        "فقط خروجی دارای Run ID، فایل ورودی، SHA-256، Version و Hash Manifest مخصوص همان Run گزارش تازه است. فایل تاریخی نباید گزارش اجرای جدید معرفی شود.",
        LIGHT_FILL,
    )

    doc.core_properties.title = "دستورالعمل اجرایی Master Baseline V3"
    doc.core_properties.subject = "Options Analytics System 3.2.0"
    doc.core_properties.author = "Options Analytics Governance"
    doc.core_properties.keywords = "Master Project Book, V3, Options, Audit, Evidence"
    doc.save(output)
    return output


def main():
    project_root = Path(__file__).resolve().parents[1]
    print(build(project_root))


if __name__ == "__main__":
    main()
